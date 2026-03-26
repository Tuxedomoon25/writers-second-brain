"""
Convert The Silence of the Bells manuscript from Obsidian markdown to
Atticus-ready .docx format.

Atticus requirements:
- Chapter titles as Heading 1 (Word default style)
- Subheadings as Heading 2 (Word default style)
- Scene breaks as *** (plain text, Normal style, left-aligned)
- No title page, TOC, or copyright page
- No images embedded
- No custom styles
- .docx format
"""

import re
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

VAULT = r"C:\Users\Timo\Obsidian Vaults\The Silence of the Bells\1. Manuscript"
OUTPUT = r"C:\Users\Timo\Desktop\The Silence of the Bells - Atticus Import.docx"

# Chapter order
CHAPTERS = [
    (r"Act 1\Chapter 1.md", "Chapter 1"),
    (r"Act 1\Chapter 2.md", "Chapter 2"),
    (r"Act 1\Chapter 3.md", "Chapter 3"),
    (r"Act 2\Chapter 4.md", "Chapter 4"),
    (r"Act 2\Chapter 5.md", "Chapter 5"),
    (r"Act 2\Chapter 6.md", "Chapter 6"),
    (r"Act 2\Chapter 7.md", "Chapter 7"),
    (r"Act 2\Chapter 8.md", "Chapter 8"),
    (r"Act 3\Chapter 9.md", "Chapter 9"),
    (r"Act 3\Chapter 10.md", "Chapter 10"),
    (r"Act 3\Chapter 11.md", "Chapter 11"),
    (r"Act 3\Chapter 12.md", "Chapter 12"),
]


def clean_scene_title(title: str) -> str:
    """Remove (POLISHED), (POLISHED — ...), version tags, etc."""
    title = re.sub(r'\s*\(POLISHED[^)]*\)', '', title)
    title = re.sub(r'\s*\(v\d[^)]*\)', '', title)
    return title.strip()


def is_metadata_line(line: str) -> bool:
    """Check if a line starts a metadata/notes section we want to skip."""
    metadata_markers = [
        '**Word Count**',
        '**Quality Scores**',
        '## Changes Summary',
        '**Mythology Enrichment**',
        '**Polish Fixes Applied**',
        '**Major Revisions**',
        '**Voice Adjustments**',
        '**Sensory Enhancements**',
        '**Negative Construction Fixes**',
        '**Rule 11 Compliance**',
        '**Rule 11',
        '**Previous Changes',
        '**Two-Layer moments**',
        '**Somatic tells used',
        '**Chapter',
        '**Ch5 somatic',
        '**Ch4 somatic',
        '### Act 3',
        '### Somatic',
        '### Quality',
        '### Changes',
        '### Act 3 Budget',
        '**Mythology',
        '### NOVEL STATUS',
        '**Estimated total',
        '**Chapter 9',
        '**Chapter 10',
        '**Chapter 11',
        '**Chapter 12',
    ]
    stripped = line.strip()
    # Catch "Word Count:" or "**Word Count" variations
    if re.match(r'\*?\*?Word Count\*?\*?\s*:', stripped):
        return True
    for marker in metadata_markers:
        if stripped.startswith(marker):
            return True
    return False


def extract_prose(filepath: str) -> list:
    """
    Extract prose content from a chapter markdown file.
    Returns a list of elements: ('heading', text), ('scene_break',), ('paragraph', text)
    """
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    elements = []
    in_metadata = False
    prev_was_blank = False
    seen_headings = set()  # track duplicate headings

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip YAML frontmatter
        if i == 0 and stripped == '---':
            i += 1
            while i < len(lines) and lines[i].strip() != '---':
                i += 1
            i += 1
            continue

        # Skip chapter-level heading (# CHAPTER X) — we add our own
        if stripped.startswith('# ') and 'CHAPTER' in stripped.upper():
            i += 1
            continue

        # Check for metadata sections
        if is_metadata_line(stripped):
            in_metadata = True
            i += 1
            continue

        # A horizontal rule (---) can either be a scene break or end of metadata
        if stripped == '---':
            if in_metadata:
                in_metadata = False
                i += 1
                continue
            else:
                # Check if this --- precedes metadata (look ahead)
                # If the next non-blank line is metadata, skip the ---
                j = i + 1
                while j < len(lines) and lines[j].strip() == '':
                    j += 1
                if j < len(lines) and is_metadata_line(lines[j].strip()):
                    in_metadata = True
                    i += 1
                    continue

                # Check if next non-blank line is a ## heading (scene break + new scene)
                if j < len(lines) and lines[j].strip().startswith('## '):
                    # This is a scene break before a new scene heading
                    elements.append(('scene_break',))
                    i += 1
                    continue

                # Otherwise it's a scene break within a scene
                elements.append(('scene_break',))
                i += 1
                continue

        if in_metadata:
            # Check if we've left metadata (blank line followed by ## heading or prose)
            if stripped == '':
                # Look ahead for next non-blank
                j = i + 1
                while j < len(lines) and lines[j].strip() == '':
                    j += 1
                if j < len(lines):
                    next_stripped = lines[j].strip()
                    if next_stripped.startswith('## ') and not is_metadata_line(next_stripped):
                        in_metadata = False
                    elif next_stripped == '---':
                        # Could be end of metadata block
                        pass
                    elif not is_metadata_line(next_stripped) and not next_stripped.startswith('-') and not next_stripped.startswith('|') and not next_stripped.startswith('*'):
                        # Looks like prose, exit metadata
                        in_metadata = False
            i += 1
            continue

        # Skip blank lines (track them for paragraph spacing)
        if stripped == '':
            prev_was_blank = True
            i += 1
            continue

        # Skip lines that look like metadata tables
        if stripped.startswith('|') and '|' in stripped[1:]:
            i += 1
            continue

        # Skip bullet points in metadata
        if stripped.startswith('- ') and any(kw in stripped for kw in [
            'threshold', 'PASS', 'adverb', 'metaphor', 'Kill', 'Remove',
            'Convert', 'Enhance', 'Strengthen', 'Adjust', 'Tighten',
            'Fix', 'Retain', 'Changed', 'Scene 1', 'Scene 2', 'Scene 3',
            'Sc1', 'Sc2', 'Sc3', 'Drift', 'Mass', 'Un-naming',
            'Cole', 'Chest', 'Root', 'Dad', 'Door', 'Postcard',
            'Hiss', 'Gate', 'Stairs', 'Map', 'Maya\'s',
            'Five-stage', 'Temperature', 'Cobblestones', 'Final line',
            'Somatic tell', 'NO copper', 'No jaw', 'Tell ARC',
            'No copper',
        ]):
            i += 1
            continue

        # Scene heading (## Title)
        if stripped.startswith('## '):
            title = stripped[3:].strip()
            title = clean_scene_title(title)

            # Skip duplicate headings (some scenes have both raw and POLISHED versions)
            if title in seen_headings:
                i += 1
                continue
            seen_headings.add(title)

            elements.append(('heading', title))
            prev_was_blank = False
            i += 1
            continue

        # Regular paragraph
        elements.append(('paragraph', stripped))
        prev_was_blank = False
        i += 1

    return elements


def add_formatted_text(paragraph, text: str):
    """
    Parse markdown inline formatting (*italic*, **bold**, ***bold italic***)
    and add runs to the paragraph.
    """
    # Pattern to match ***bold italic***, **bold**, *italic*, _italic_
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|_(.+?)_)'

    pos = 0
    for match in re.finditer(pattern, text):
        # Add plain text before this match
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])

        if match.group(2):  # ***bold italic***
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(3):  # **bold**
            run = paragraph.add_run(match.group(3))
            run.bold = True
        elif match.group(4):  # *italic*
            run = paragraph.add_run(match.group(4))
            run.italic = True
        elif match.group(5):  # _italic_
            run = paragraph.add_run(match.group(5))
            run.italic = True

        pos = match.end()

    # Add remaining text
    if pos < len(text):
        paragraph.add_run(text[pos:])


def build_docx():
    doc = Document()

    for ch_idx, (rel_path, chapter_name) in enumerate(CHAPTERS):
        filepath = os.path.join(VAULT, rel_path)
        elements = extract_prose(filepath)

        # Add chapter heading (Heading 1)
        doc.add_heading(chapter_name, level=1)

        # Strip leading scene breaks before first content
        while elements and elements[0][0] == 'scene_break':
            elements.pop(0)

        for elem in elements:
            if elem[0] == 'heading':
                # Scene subheading (Heading 2)
                doc.add_heading(elem[1], level=2)
            elif elem[0] == 'scene_break':
                # Atticus scene break: *** on its own line, Normal style
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.style = doc.styles['Normal']
                p.add_run('***')
            elif elem[0] == 'paragraph':
                p = doc.add_paragraph()
                p.style = doc.styles['Normal']
                add_formatted_text(p, elem[1])

        # Add page break between chapters (except after last)
        if ch_idx < len(CHAPTERS) - 1:
            doc.add_page_break()

    doc.save(OUTPUT)
    print(f"Saved to: {OUTPUT}")
    print(f"Chapters: {len(CHAPTERS)}")

    # Count words
    total_words = 0
    for para in doc.paragraphs:
        total_words += len(para.text.split())
    print(f"Total words: {total_words}")


if __name__ == '__main__':
    build_docx()
