"""
Convert nonfiction manuscript from Obsidian markdown to print-ready .docx.

Usage:
    python format_nonfiction.py --config path/to/format-config.yaml --format atticus
    python format_nonfiction.py --config path/to/format-config.yaml --format indesign
    python format_nonfiction.py --config path/to/format-config.yaml --format both

Handles nonfiction-specific features:
  - Citations: [@Author.Work, p. X] -> Word footnotes
  - Flat chapter structure (# headings as chapter boundaries)
  - Epigraphs (> blockquotes) styled as Block Quote
  - Subheadings (## or ###) as H2
"""

import argparse
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).parent))
from lib.parser import (
    load_config, extract_nonfiction_prose, add_formatted_text,
    strip_citations_for_text, extract_citation_refs,
    generate_indesign_config,
)


def convert_citation_to_footnote(paragraph, citation_text: str, refs: list, style: str = 'conversational'):
    """
    Convert citation markers to Word footnotes.

    Handles locator formats: page numbers, chapter refs, essay titles, or none.
    """
    footnote_parts = []
    for author, work, locator in refs:
        # Convert CamelCase work title to spaced
        work_spaced = re.sub(r'(?<=[a-z])(?=[A-Z])', ' ', work)
        prefix = "See " if style == 'conversational' else ""
        if locator:
            locator = locator.strip()
            # Page reference (p. 136, pp. 6, 18)
            if re.match(r'p+\.', locator):
                footnote_parts.append(f"{prefix}{author}, {work_spaced}, {locator}")
            # Essay title in quotes ("Myth Became Fact")
            elif locator.startswith('"'):
                footnote_parts.append(f"{prefix}{author}, {work_spaced}, {locator}")
            # Book/Chapter reference (Book II, Ch. 1)
            else:
                footnote_parts.append(f"{prefix}{author}, {work_spaced}, {locator}")
        else:
            footnote_parts.append(f"{prefix}{author}, {work_spaced}")

    return "; ".join(footnote_parts)


def add_paragraph_with_citations(doc, paragraph_doc, text: str, refs: list,
                                  citation_style: str = 'conversational'):
    """
    Add text to a paragraph, converting inline citations to superscript footnote markers.
    Since python-docx doesn't support real Word footnotes natively, we use
    a superscript number and collect footnotes at chapter end.
    """
    # Strip citation markers from display text
    clean_text = strip_citations_for_text(text)
    add_formatted_text(paragraph_doc, clean_text)

    # Add footnote reference as superscript
    if refs:
        footnote_text = convert_citation_to_footnote(paragraph_doc, text, refs, citation_style)
        run = paragraph_doc.add_run(f" [{footnote_text}]")
        run.font.size = Pt(8)
        run.font.superscript = True


def build_atticus_docx(config: dict) -> str:
    """Build Atticus-format .docx for nonfiction: H1 chapters, body text, epigraphs."""
    doc = Document()

    vault_path = config['vault_path']
    manuscript_path = config['manuscript_path']
    chapters = config['chapters']
    citation_style = config.get('citation_style', 'conversational')
    has_citations = config.get('has_citations', False)

    for ch_idx, chapter in enumerate(chapters):
        rel_path = chapter['path']
        chapter_name = chapter['name']
        filepath = os.path.join(vault_path, manuscript_path, rel_path)

        if not os.path.exists(filepath):
            print(f"  WARNING: File not found: {filepath}")
            continue

        # For single-file manuscripts, extract specific section
        section_name = chapter.get('section')
        elements = extract_nonfiction_prose(filepath)

        # If section specified, filter to just that section
        if section_name:
            elements = filter_to_section(elements, section_name)

        is_part = chapter.get('type') == 'part'

        if not elements and not is_part:
            print(f"  WARNING: No content found for chapter '{chapter_name}'")
            continue

        if is_part:
            # Part divider page: centered title + italic subtitle
            p = doc.add_heading(chapter_name, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for elem in elements:
                if elem[0] == 'subheading':
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(elem[1])
                    run.italic = True
        else:
            # Regular chapter
            doc.add_heading(chapter_name, level=1)

            for elem in elements:
                if elem[0] == 'chapter':
                    continue
                elif elem[0] == 'subheading':
                    doc.add_heading(elem[1], level=2)
                elif elem[0] == 'epigraph':
                    p = doc.add_paragraph()
                    p.style = doc.styles['Normal']
                    run = p.add_run(elem[1])
                    run.italic = True
                elif elem[0] == 'scene_break':
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.style = doc.styles['Normal']
                    p.add_run('***')
                elif elem[0] == 'citation' and has_citations:
                    p = doc.add_paragraph()
                    p.style = doc.styles['Normal']
                    add_paragraph_with_citations(doc, p, elem[1], extract_citation_refs(elem[1]), citation_style)
                elif elem[0] == 'list_item':
                    p = doc.add_paragraph()
                    p.style = doc.styles['List Bullet']
                    add_formatted_text(p, elem[1])
                elif elem[0] == 'paragraph':
                    p = doc.add_paragraph()
                    p.style = doc.styles['Normal']
                    add_formatted_text(p, elem[1])

        # Page break between chapters (except after last)
        if ch_idx < len(chapters) - 1:
            doc.add_page_break()

    output_dir = config['output_path']
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"{config['title']} - Atticus Import.docx")

    doc.save(output_path)

    h1_count = sum(1 for p in doc.paragraphs if p.style.name == 'Heading 1')
    h2_count = sum(1 for p in doc.paragraphs if p.style.name == 'Heading 2')
    body_paras = sum(1 for p in doc.paragraphs if p.style.name == 'Normal' and p.text.strip() and p.text.strip() != '***')

    print(f"  Atticus .docx saved: {output_path}")
    print(f"  H1 (chapters): {h1_count}")
    print(f"  H2 (subheadings): {h2_count}")
    print(f"  Body paragraphs: {body_paras}")

    return output_path


def create_custom_style(doc, name, base_style='Normal'):
    """Create a custom paragraph style."""
    styles = doc.styles
    style = styles.add_style(name, 1)
    style.base_style = styles[base_style]
    return style


def build_indesign_docx(config: dict) -> str:
    """Build InDesign-format .docx for nonfiction with custom styles."""
    doc = Document()

    style_chapter_title = create_custom_style(doc, 'ChapterTitle')
    style_part_title = create_custom_style(doc, 'PartTitle')
    style_part_subtitle = create_custom_style(doc, 'PartSubtitle')
    style_scene_break = create_custom_style(doc, 'SceneBreak')
    style_epigraph = create_custom_style(doc, 'Epigraph')

    vault_path = config['vault_path']
    manuscript_path = config['manuscript_path']
    chapters = config['chapters']
    citation_style = config.get('citation_style', 'conversational')
    has_citations = config.get('has_citations', False)

    for ch_idx, chapter in enumerate(chapters):
        rel_path = chapter['path']
        chapter_name = chapter['name']
        filepath = os.path.join(vault_path, manuscript_path, rel_path)

        if not os.path.exists(filepath):
            print(f"  WARNING: File not found: {filepath}")
            continue

        section_name = chapter.get('section')
        elements = extract_nonfiction_prose(filepath)

        if section_name:
            elements = filter_to_section(elements, section_name)

        is_part = chapter.get('type') == 'part'

        if not elements and not is_part:
            print(f"  WARNING: No content found for chapter '{chapter_name}'")
            continue

        if is_part:
            # Part divider page: PartTitle + PartSubtitle styles
            p = doc.add_paragraph()
            p.style = style_part_title
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(chapter_name)
            for elem in elements:
                if elem[0] == 'subheading':
                    p = doc.add_paragraph()
                    p.style = style_part_subtitle
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(elem[1])
                    run.italic = True
        else:
            # Regular chapter
            p = doc.add_paragraph()
            p.style = style_chapter_title
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(chapter_name)

            for elem in elements:
                if elem[0] == 'chapter':
                    continue
                elif elem[0] == 'subheading':
                    doc.add_heading(elem[1], level=2)
                elif elem[0] == 'epigraph':
                    p = doc.add_paragraph()
                    p.style = style_epigraph
                    run = p.add_run(elem[1])
                    run.italic = True
                elif elem[0] == 'scene_break':
                    p = doc.add_paragraph()
                    p.style = style_scene_break
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run('***')
                elif elem[0] == 'citation' and has_citations:
                    p = doc.add_paragraph()
                    p.style = doc.styles['Normal']
                    add_paragraph_with_citations(doc, p, elem[1], extract_citation_refs(elem[1]), citation_style)
                elif elem[0] == 'list_item':
                    p = doc.add_paragraph()
                    p.style = doc.styles['Normal']
                    add_formatted_text(p, elem[1])
                elif elem[0] == 'paragraph':
                    p = doc.add_paragraph()
                    p.style = doc.styles['Normal']
                    add_formatted_text(p, elem[1])

        if ch_idx < len(chapters) - 1:
            doc.add_page_break()

    output_dir = config['output_path']
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"{config['title']} - InDesign Import.docx")

    doc.save(output_path)

    chapter_titles = sum(1 for p in doc.paragraphs if p.style.name == 'ChapterTitle')
    total_words = sum(len(p.text.split()) for p in doc.paragraphs if p.style.name == 'Normal')

    print(f"  InDesign .docx saved: {output_path}")
    print(f"  Chapter titles: {chapter_titles}")
    print(f"  Prose words: {total_words}")

    return output_path


def filter_to_section(elements: list, section_name: str) -> list:
    """
    Filter elements to a specific section (for single-file manuscripts).
    Extracts content between the matching # heading and the next # heading.
    """
    filtered = []
    in_section = False

    for elem in elements:
        if elem[0] == 'chapter':
            if elem[1].upper() == section_name.upper():
                in_section = True
                continue  # Don't include the chapter marker itself
            elif in_section:
                break  # Reached the next chapter
        elif in_section:
            filtered.append(elem)

    return filtered


def main():
    parser = argparse.ArgumentParser(description='Convert nonfiction manuscript to print-ready .docx')
    parser.add_argument('--config', required=True, help='Path to format-config.yaml')
    parser.add_argument('--format', required=True, choices=['atticus', 'indesign', 'both'],
                        help='Output format: atticus, indesign, or both')
    args = parser.parse_args()

    config = load_config(args.config)

    print(f'Formatting: {config["title"]}')
    print(f'Chapters: {len(config["chapters"])}')
    print(f'Structure: {config["structure"]}')
    print(f'Citations: {config.get("has_citations", False)}')
    print(f'Output: {config["output_path"]}')
    print()

    if args.format in ('atticus', 'both'):
        print('[Atticus]')
        build_atticus_docx(config)
        print()

    if args.format in ('indesign', 'both'):
        print('[InDesign]')
        indesign_path = build_indesign_docx(config)
        print()

        print('[InDesign Config]')
        config_json_path = generate_indesign_config(config, config['output_path'])
        print(f'  Config JSON saved: {config_json_path}')
        print()

    print('Done.')


if __name__ == '__main__':
    main()
