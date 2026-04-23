"""
Convert fiction manuscript from Obsidian markdown to print-ready .docx.

Usage:
    python format_fiction.py --config path/to/format-config.yaml --format atticus
    python format_fiction.py --config path/to/format-config.yaml --format indesign
    python format_fiction.py --config path/to/format-config.yaml --format both

Replaces the three hardcoded scripts:
  - convert_to_docx.py
  - convert_to_docx_atticus.py
  - convert_to_docx_indesign.py
"""

import argparse
import os
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add parent dir to path for lib imports
sys.path.insert(0, str(Path(__file__).parent))
from lib.parser import (
    load_config, extract_fiction_prose, add_formatted_text,
    generate_indesign_config,
)


def _atticus_centered_paragraph(doc, text, italic=False, bold=False, size=None):
    """Add a centered paragraph (helper for Atticus front/back matter)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    if italic:
        run.italic = True
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    return p


def _atticus_body_paragraph(doc, text, italic=False, center=False):
    """Add a body paragraph (helper for Atticus front/back matter)."""
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    if italic:
        run.italic = True
    return p


def _atticus_add_front_matter(doc, config):
    """Add front matter pages: title, copyright, dedication."""
    title = config.get('title', '')
    subtitle = config.get('subtitle', '')
    author = config.get('author', '')
    website = config.get('website', '')
    isbn = config.get('isbn', '')
    dedication = config.get('dedication', '')
    copyright_text = config.get('copyright', '')

    # Title page
    _atticus_centered_paragraph(doc, '', size=24)
    _atticus_centered_paragraph(doc, '', size=24)
    _atticus_centered_paragraph(doc, title, bold=True, size=28)
    if subtitle:
        _atticus_centered_paragraph(doc, '', size=12)
        _atticus_centered_paragraph(doc, subtitle, italic=True, size=14)
    _atticus_centered_paragraph(doc, '', size=12)
    _atticus_centered_paragraph(doc, '', size=12)
    _atticus_centered_paragraph(doc, author, size=14)
    doc.add_page_break()

    # Copyright page
    if copyright_text:
        for line in copyright_text.split('\n'):
            line = line.strip()
            if line:
                p = doc.add_paragraph()
                p.style = doc.styles['Normal']
                run = p.add_run(line)
                run.font.size = Pt(9)
            else:
                doc.add_paragraph()
        if website:
            p = doc.add_paragraph()
            run = p.add_run(website)
            run.font.size = Pt(9)
        if isbn:
            p = doc.add_paragraph()
            run = p.add_run(f'ISBN: {isbn}')
            run.font.size = Pt(9)
        doc.add_page_break()

    # Dedication page
    if dedication:
        _atticus_centered_paragraph(doc, '', size=12)
        _atticus_centered_paragraph(doc, '', size=12)
        _atticus_centered_paragraph(doc, '', size=12)
        for line in dedication.split('\n'):
            if line.strip():
                _atticus_centered_paragraph(doc, line.strip(), italic=True, size=14)
        doc.add_page_break()


def _atticus_add_back_matter(doc, config):
    """Add back matter pages: epilogue, about author, newsletter, coming next, also by."""
    epilogue = config.get('epilogue', '')
    bm = config.get('back_matter', {}) or {}

    # Epilogue
    if epilogue:
        doc.add_page_break()
        doc.add_heading('Epilogue', level=1)
        for line in epilogue.split('\n'):
            line = line.strip()
            if line:
                _atticus_body_paragraph(doc, line)

    # About the Author
    about = bm.get('about_author', '')
    if about:
        doc.add_page_break()
        doc.add_heading('About the Author', level=1)
        for line in about.split('\n'):
            line = line.strip()
            if line:
                _atticus_body_paragraph(doc, line)

    # Newsletter / Stay in the Sung World
    nl = bm.get('newsletter')
    if nl:
        doc.add_page_break()
        if nl.get('heading'):
            doc.add_heading(nl['heading'], level=1)
        if nl.get('body'):
            for line in nl['body'].split('\n'):
                line = line.strip()
                if line:
                    _atticus_body_paragraph(doc, line, italic=True, center=True)
        if nl.get('url'):
            _atticus_centered_paragraph(doc, '', size=12)
            _atticus_centered_paragraph(doc, nl['url'], bold=True, size=12)

    # Coming Next
    cn = bm.get('coming_next')
    if cn:
        doc.add_page_break()
        if cn.get('heading'):
            doc.add_heading(cn['heading'], level=1)
        if cn.get('body'):
            for line in cn['body'].split('\n'):
                line = line.strip()
                if line:
                    _atticus_body_paragraph(doc, line, italic=True, center=True)
        _atticus_centered_paragraph(doc, '', size=12)
        if cn.get('book_label'):
            _atticus_centered_paragraph(doc, cn['book_label'], bold=True, size=14)
        if cn.get('series_label'):
            _atticus_centered_paragraph(doc, cn['series_label'], size=11)
        if cn.get('coming_label'):
            _atticus_centered_paragraph(doc, cn['coming_label'], italic=True, size=11)

    # Also By
    ab = bm.get('also_by')
    if ab:
        doc.add_page_break()
        if ab.get('heading'):
            doc.add_heading(ab['heading'], level=1)
        if ab.get('cycle_label'):
            _atticus_centered_paragraph(doc, '', size=10)
            _atticus_centered_paragraph(doc, ab['cycle_label'], bold=True, size=11)
        for title in ab.get('cycle_titles', []) or []:
            _atticus_centered_paragraph(doc, title, italic=True, size=13)
        if ab.get('cycle_subline'):
            _atticus_centered_paragraph(doc, ab['cycle_subline'], size=10)
        if ab.get('cycle_tagline'):
            _atticus_centered_paragraph(doc, '', size=10)
            for line in ab['cycle_tagline'].split('\n'):
                line = line.strip()
                if line:
                    _atticus_centered_paragraph(doc, line, italic=True, size=11)
        if ab.get('nonfiction_label'):
            _atticus_centered_paragraph(doc, '', size=10)
            _atticus_centered_paragraph(doc, ab['nonfiction_label'], bold=True, size=11)
        for title in ab.get('nonfiction_titles', []) or []:
            _atticus_centered_paragraph(doc, title, italic=True, size=13)
        if ab.get('closing_tagline'):
            _atticus_centered_paragraph(doc, '', size=10)
            _atticus_centered_paragraph(doc, '', size=10)
            for line in ab['closing_tagline'].split('\n'):
                line = line.strip()
                if line:
                    _atticus_centered_paragraph(doc, line, italic=True, size=11)
        if ab.get('url'):
            _atticus_centered_paragraph(doc, '', size=10)
            _atticus_centered_paragraph(doc, ab['url'], bold=True, size=12)


def build_atticus_docx(config: dict) -> str:
    """
    Build Atticus-format .docx: H1 chapter names, H2 scene names, *** scene breaks.
    Includes front matter (title, copyright, dedication) and back matter
    (epilogue, about author, newsletter, coming next, also by) from the config.
    Returns the output path.
    """
    doc = Document()

    # Front matter
    _atticus_add_front_matter(doc, config)

    vault_path = config['vault_path']
    manuscript_path = config['manuscript_path']
    chapters = config['chapters']

    for ch_idx, chapter in enumerate(chapters):
        rel_path = chapter['path']
        chapter_name = chapter['name']
        filepath = os.path.join(vault_path, manuscript_path, rel_path)

        if not os.path.exists(filepath):
            print(f"  WARNING: Chapter file not found: {filepath}")
            continue

        elements = extract_fiction_prose(filepath)

        # Chapter heading (H1) = chapter NAME
        doc.add_heading(chapter_name, level=1)

        # Strip leading scene breaks before first content
        while elements and elements[0][0] == 'scene_break':
            elements.pop(0)

        for elem in elements:
            if elem[0] == 'heading':
                # Scene headings removed — chapter name is sufficient
                continue
            elif elem[0] == 'scene_break':
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.style = doc.styles['Normal']
                p.add_run('***')
            elif elem[0] == 'paragraph':
                p = doc.add_paragraph()
                p.style = doc.styles['Normal']
                add_formatted_text(p, elem[1])

        # Page break between chapters (except after last)
        if ch_idx < len(chapters) - 1:
            doc.add_page_break()

    # Back matter
    _atticus_add_back_matter(doc, config)

    # Build output path
    title_slug = config['title'].replace(' ', '_')
    output_dir = config['output_path']
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"{config['title']} - Atticus Import.docx")

    doc.save(output_path)

    # Stats
    h1_count = sum(1 for p in doc.paragraphs if p.style.name == 'Heading 1')
    h2_count = sum(1 for p in doc.paragraphs if p.style.name == 'Heading 2')
    scene_breaks = sum(1 for p in doc.paragraphs if p.text.strip() == '***')
    body_paras = sum(1 for p in doc.paragraphs if p.style.name == 'Normal' and p.text.strip() and p.text.strip() != '***')

    print(f"  Atticus .docx saved: {output_path}")
    print(f"  H1 (chapters): {h1_count}")
    print(f"  H2 (scenes): {h2_count}")
    print(f"  Scene breaks (***): {scene_breaks}")
    print(f"  Body paragraphs: {body_paras}")

    return output_path


def create_custom_style(doc, name, base_style='Normal'):
    """Create a custom paragraph style in the Word document."""
    styles = doc.styles
    style = styles.add_style(name, 1)  # 1 = WD_STYLE_TYPE.PARAGRAPH
    style.base_style = styles[base_style]
    return style


def build_indesign_docx(config: dict) -> str:
    """
    Build InDesign-format .docx: ChapterNumber + ChapterTitle custom styles,
    scene titles removed, SceneBreak custom style.
    Returns the output path.
    """
    doc = Document()

    # Create custom Word styles that InDesign will map from
    style_chapter_num = create_custom_style(doc, 'ChapterNumber')
    style_chapter_title = create_custom_style(doc, 'ChapterTitle')
    style_scene_break = create_custom_style(doc, 'SceneBreak')

    vault_path = config['vault_path']
    manuscript_path = config['manuscript_path']
    chapters = config['chapters']

    for ch_idx, chapter in enumerate(chapters):
        rel_path = chapter['path']
        ch_number = chapter.get('number', ch_idx + 1)
        ch_name = chapter['name']
        filepath = os.path.join(vault_path, manuscript_path, rel_path)

        if not os.path.exists(filepath):
            print(f"  WARNING: Chapter file not found: {filepath}")
            continue

        elements = extract_fiction_prose(filepath)

        # Strip leading scene breaks
        while elements and elements[0][0] == 'scene_break':
            elements.pop(0)

        # Add Chapter Number paragraph (just the digit)
        p = doc.add_paragraph()
        p.style = style_chapter_num
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(str(ch_number))

        # Add Chapter Title paragraph (chapter NAME)
        p = doc.add_paragraph()
        p.style = style_chapter_title
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(ch_name)

        # Collapse consecutive scene_break + heading pairs into single scene_break
        # (headings are skipped for print, so scene_break before heading is redundant)
        filtered = []
        for elem in elements:
            if elem[0] == 'heading':
                continue  # Scene headings REMOVED for print
            if elem[0] == 'scene_break' and filtered and filtered[-1][0] == 'scene_break':
                continue  # Skip duplicate consecutive scene breaks
            filtered.append(elem)

        for elem in filtered:
            if elem[0] == 'scene_break':
                # Whitespace-only paragraph to trigger drop cap on next paragraph
                p = doc.add_paragraph()
                p.style = style_scene_break
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run('\u2009')  # thin space — gives paragraph height in InDesign
            elif elem[0] == 'paragraph':
                p = doc.add_paragraph()
                p.style = doc.styles['Normal']
                add_formatted_text(p, elem[1])

        # Page break between chapters (except after last)
        if ch_idx < len(chapters) - 1:
            doc.add_page_break()

    # Build output path
    output_dir = config['output_path']
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"{config['title']} - InDesign Import.docx")

    doc.save(output_path)

    # Stats
    chapter_nums = sum(1 for p in doc.paragraphs if p.style.name == 'ChapterNumber')
    chapter_titles = sum(1 for p in doc.paragraphs if p.style.name == 'ChapterTitle')
    scene_breaks = sum(1 for p in doc.paragraphs if p.style.name == 'SceneBreak')
    total_words = sum(len(p.text.split()) for p in doc.paragraphs if p.style.name == 'Normal')

    print(f"  InDesign .docx saved: {output_path}")
    print(f"  Chapter numbers: {chapter_nums}")
    print(f"  Chapter titles: {chapter_titles}")
    print(f"  Scene breaks: {scene_breaks}")
    print(f"  Prose words: {total_words}")

    return output_path


def main():
    parser = argparse.ArgumentParser(description='Convert fiction manuscript to print-ready .docx')
    parser.add_argument('--config', required=True, help='Path to format-config.yaml')
    parser.add_argument('--format', required=True, choices=['atticus', 'indesign', 'both'],
                        help='Output format: atticus, indesign, or both')
    args = parser.parse_args()

    config = load_config(args.config)

    print(f'Formatting: {config["title"]}')
    print(f'Chapters: {len(config["chapters"])}')
    print(f'Structure: {config["structure"]}')
    print(f'Output: {config["output_path"]}')
    print()

    if args.format in ('atticus', 'both'):
        print('[Atticus]')
        build_atticus_docx(config)
        print()

    if args.format in ('indesign', 'both'):
        print('[InDesign]')
        indesign_path = build_indesign_docx(config)
        print()

        # Generate InDesign config.json for JSX scripts
        print('[InDesign Config]')
        config_json_path = generate_indesign_config(config, config['output_path'])
        print(f'  Config JSON saved: {config_json_path}')
        print()

    print('Done.')


if __name__ == '__main__':
    main()
