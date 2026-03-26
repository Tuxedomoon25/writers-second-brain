"""
Convert The Silence of the Bells manuscript to Atticus-ready .docx.

Atticus requirements:
- Chapter titles as Heading 1 (Word default style) — uses chapter NAMES
- Scene titles as Heading 2 (Word default style)
- Scene breaks as *** ONLY between scenes (not within scenes)
- No title page, TOC, or copyright page
- No custom styles — only Word defaults
- .docx format
"""

import re
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

VAULT = r"C:\Users\Timo\Obsidian Vaults\The Silence of the Bells\1. Manuscript"
OUTPUT = r"C:\Users\Timo\Desktop\Projects\Atticus Indesign\The Silence of the Bells - Atticus Import.docx"

# Chapter order: (file path, chapter name)
CHAPTERS = [
    (r"Act 1\Chapter 1.md", "The Grey City"),
    (r"Act 1\Chapter 2.md", "The Impossible Door"),
    (r"Act 1\Chapter 3.md", "The Fall"),
    (r"Act 2\Chapter 4.md", "The Fortress of Last Words"),
    (r"Act 2\Chapter 5.md", "The Engineer"),
    (r"Act 2\Chapter 6.md", "The Broken Mirror"),
    (r"Act 2\Chapter 7.md", "The Gilded Cage"),
    (r"Act 2\Chapter 8.md", "The Ascent"),
    (r"Act 3\Chapter 9.md", "The Hollow Machine"),
    (r"Act 3\Chapter 10.md", "The Engineer's End"),
    (r"Act 3\Chapter 11.md", "The Living Clapper"),
    (r"Act 3\Chapter 12.md", "The Golden Echo"),
]


def clean_scene_title(title: str) -> str:
    title = re.sub(r'\s*\(POLISHED[^)]*\)', '', title)
    title = re.sub(r'\s*\(v\d[^)]*\)', '', title)
    return title.strip()


def is_metadata_line(line: str) -> bool:
    metadata_markers = [
        '**Word Count**', '**Quality Scores**', '## Changes Summary',
        '**Mythology Enrichment**', '**Polish Fixes Applied**',
        '**Major Revisions**', '**Voice Adjustments**',
        '**Sensory Enhancements**', '**Negative Construction Fixes**',
        '**Rule 11 Compliance**', '**Rule 11', '**Previous Changes',
        '**Two-Layer moments**', '**Somatic tells used',
        '**Chapter', '**Ch5 somatic', '**Ch4 somatic',
        '### Act 3', '### Somatic', '### Quality', '### Changes',
        '### Act 3 Budget', '**Mythology', '### NOVEL STATUS',
        '**Estimated total', '**Chapter 9', '**Chapter 10',
        '**Chapter 11', '**Chapter 12',
    ]
    stripped = line.strip()
    if re.match(r'\*?\*?Word Count\*?\*?\s*:', stripped):
        return True
    for marker in metadata_markers:
        if stripped.startswith(marker):
            return True
    return False


def extract_prose(filepath: str) -> list:
    """
    Extract prose content from a chapter markdown file.
    Returns list of: ('heading', text), ('scene_break',), ('paragraph', text)

    IMPORTANT: Only emits scene_break BETWEEN scenes (before a ## heading),
    not for internal --- breaks within a scene.
    """
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    elements = []
    in_metadata = False
    seen_headings = set()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip YAML frontmatter
        if i == 0 and stripped == '---':
            i += 1
            while i < len(lines) and lines[i].strip() != '---':
                i += 1
            i += 1
            continue

        # Skip chapter-level heading (# CHAPTER X)
        if stripped.startswith('# ') and 'CHAPTER' in stripped.upper():
            i += 1
            continue

        # Check for metadata sections
        if is_metadata_line(stripped):
            in_metadata = True
            i += 1
            continue

        # Handle horizontal rules (---)
        if stripped == '---':
            if in_metadata:
                in_metadata = False
                i += 1
                continue

            # Look ahead: does this --- precede metadata?
            j = i + 1
            while j < len(lines) and lines[j].strip() == '':
                j += 1
            if j < len(lines) and is_metadata_line(lines[j].strip()):
                in_metadata = True
                i += 1
                continue

            # Does this --- precede a ## heading? = scene break
            if j < len(lines) and lines[j].strip().startswith('## '):
                # This is a scene transition — emit scene_break
                elements.append(('scene_break',))
                i += 1
                continue

            # Otherwise this is an INTERNAL break within a scene
            # For Atticus: also emit as scene break (*** between sub-sections)
            # But the user wants only 3 scenes per chapter, so we SKIP internal breaks
            # and let the prose flow continuously within each scene
            i += 1
            continue

        if in_metadata:
            if stripped == '':
                j = i + 1
                while j < len(lines) and lines[j].strip() == '':
                    j += 1
                if j < len(lines):
                    next_stripped = lines[j].strip()
                    if next_stripped.startswith('## ') and not is_metadata_line(next_stripped):
                        in_metadata = False
                    elif next_stripped == '---':
                        pass
                    elif not is_metadata_line(next_stripped) and not next_stripped.startswith('-') and not next_stripped.startswith('|') and not next_stripped.startswith('*'):
                        in_metadata = False
            i += 1
            continue

        if stripped == '':
            i += 1
            continue

        if stripped.startswith('|') and '|' in stripped[1:]:
            i += 1
            continue

        if stripped.startswith('- ') and any(kw in stripped for kw in [
            'threshold', 'PASS', 'adverb', 'metaphor', 'Kill', 'Remove',
            'Convert', 'Enhance', 'Strengthen', 'Adjust', 'Tighten',
            'Fix', 'Retain', 'Changed', 'Scene 1', 'Scene 2', 'Scene 3',
            'Sc1', 'Sc2', 'Sc3', 'Drift', 'Mass', 'Un-naming',
            'Cole', 'Chest', 'Root', 'Dad', 'Door', 'Postcard',
            'Hiss', 'Gate', 'Stairs', 'Map', 'Maya\'s',
            'Five-stage', 'Temperature', 'Cobblestones', 'Final line',
            'Somatic tell', 'NO copper', 'No jaw', 'Tell ARC',
            'No copper',
        ]):
            i += 1
            continue

        # Scene heading (## Title)
        if stripped.startswith('## '):
            title = stripped[3:].strip()
            title = clean_scene_title(title)
            if title in seen_headings:
                i += 1
                continue
            seen_headings.add(title)
            elements.append(('heading', title))
            i += 1
            continue

        # Regular paragraph
        elements.append(('paragraph', stripped))
        i += 1

    # Post-process: ensure *** between every scene heading (except first)
    fixed = []
    first_heading_seen = False
    for elem in elements:
        if elem[0] == 'heading':
            if first_heading_seen:
                # Check if previous element is already a scene_break
                if not fixed or fixed[-1][0] != 'scene_break':
                    fixed.append(('scene_break',))
            first_heading_seen = True
        fixed.append(elem)

    return fixed


def add_formatted_text(paragraph, text: str):
    """Parse markdown inline formatting and add runs."""
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|_(.+?)_)'
    pos = 0
    for match in re.finditer(pattern, text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        if match.group(2):
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(3):
            run = paragraph.add_run(match.group(3))
            run.bold = True
        elif match.group(4):
            run = paragraph.add_run(match.group(4))
            run.italic = True
        elif match.group(5):
            run = paragraph.add_run(match.group(5))
            run.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def build_docx():
    doc = Document()

    for ch_idx, (rel_path, chapter_name) in enumerate(CHAPTERS):
        filepath = os.path.join(VAULT, rel_path)
        elements = extract_prose(filepath)

        # Chapter heading (H1) = chapter NAME
        doc.add_heading(chapter_name, level=1)

        # Strip leading scene breaks before first content
        while elements and elements[0][0] == 'scene_break':
            elements.pop(0)

        for elem in elements:
            if elem[0] == 'heading':
                # Scene name (H2)
                doc.add_heading(elem[1], level=2)
            elif elem[0] == 'scene_break':
                # *** scene break — Normal style, left-aligned, no extra formatting
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.style = doc.styles['Normal']
                p.add_run('***')
            elif elem[0] == 'paragraph':
                p = doc.add_paragraph()
                p.style = doc.styles['Normal']
                add_formatted_text(p, elem[1])

        # Page break between chapters (except after last)
        if ch_idx < len(CHAPTERS) - 1:
            doc.add_page_break()

    doc.save(OUTPUT)
    print(f"Saved to: {OUTPUT}")
    print(f"Chapters: {len(CHAPTERS)}")

    # Stats
    h1_count = 0
    h2_count = 0
    scene_breaks = 0
    body_paras = 0
    for para in doc.paragraphs:
        if para.style.name == 'Heading 1':
            h1_count += 1
        elif para.style.name == 'Heading 2':
            h2_count += 1
        elif para.text.strip() == '***':
            scene_breaks += 1
        elif para.style.name == 'Normal' and para.text.strip():
            body_paras += 1

    print(f"H1 (chapters): {h1_count}")
    print(f"H2 (scenes): {h2_count}")
    print(f"Scene breaks (***): {scene_breaks}")
    print(f"Body paragraphs: {body_paras}")
    print(f"Expected: 12 chapters, 36 scenes, 24 scene breaks (2 per chapter)")


if __name__ == '__main__':
    build_docx()
