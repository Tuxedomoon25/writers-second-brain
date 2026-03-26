"""
Convert The Silence of the Bells manuscript from Obsidian markdown to
InDesign-optimized .docx format.

InDesign requirements:
- "ChapterNumber" custom Word style for chapter digits (maps to InDesign "Chapter Number")
- "ChapterTitle" custom Word style for chapter title (maps to InDesign "Chapter Title")
- Scene titles (H2) REMOVED — only *** scene breaks between scenes
- First scene name per chapter becomes the Chapter Title
- "SceneBreak" custom Word style for *** paragraphs (maps to InDesign "Scene Break")
- Body text as Normal (maps to InDesign "Body Text")
- Italic/bold formatting preserved via runs
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

VAULT = r"C:\Users\Timo\Obsidian Vaults\The Silence of the Bells\1. Manuscript"
OUTPUT = r"C:\Users\Timo\Desktop\Projects\Atticus Indesign\The Silence of the Bells - InDesign Import.docx"

# Chapter order: (file path, chapter number, chapter name)
CHAPTERS = [
    (r"Act 1\Chapter 1.md", 1, "The Grey City"),
    (r"Act 1\Chapter 2.md", 2, "The Impossible Door"),
    (r"Act 1\Chapter 3.md", 3, "The Fall"),
    (r"Act 2\Chapter 4.md", 4, "The Fortress of Last Words"),
    (r"Act 2\Chapter 5.md", 5, "The Engineer"),
    (r"Act 2\Chapter 6.md", 6, "The Broken Mirror"),
    (r"Act 2\Chapter 7.md", 7, "The Gilded Cage"),
    (r"Act 2\Chapter 8.md", 8, "The Ascent"),
    (r"Act 3\Chapter 9.md", 9, "The Hollow Machine"),
    (r"Act 3\Chapter 10.md", 10, "The Engineer's End"),
    (r"Act 3\Chapter 11.md", 11, "The Living Clapper"),
    (r"Act 3\Chapter 12.md", 12, "The Golden Echo"),
]


# --- Reuse parsing functions from convert_to_docx.py ---

def clean_scene_title(title: str) -> str:
    title = re.sub(r'\s*\(POLISHED[^)]*\)', '', title)
    title = re.sub(r'\s*\(v\d[^)]*\)', '', title)
    return title.strip()


def is_metadata_line(line: str) -> bool:
    metadata_markers = [
        '**Word Count**', '**Quality Scores**', '## Changes Summary',
        '**Mythology Enrichment**', '**Polish Fixes Applied**',
        '**Major Revisions**', '**Voice Adjustments**',
        '**Sensory Enhancements**', '**Negative Construction Fixes**',
        '**Rule 11 Compliance**', '**Rule 11', '**Previous Changes',
        '**Two-Layer moments**', '**Somatic tells used',
        '**Chapter', '**Ch5 somatic', '**Ch4 somatic',
        '### Act 3', '### Somatic', '### Quality', '### Changes',
        '### Act 3 Budget', '**Mythology', '### NOVEL STATUS',
        '**Estimated total', '**Chapter 9', '**Chapter 10',
        '**Chapter 11', '**Chapter 12',
    ]
    stripped = line.strip()
    if re.match(r'\*?\*?Word Count\*?\*?\s*:', stripped):
        return True
    for marker in metadata_markers:
        if stripped.startswith(marker):
            return True
    return False


def extract_prose(filepath: str) -> list:
    """
    Extract prose content from a chapter markdown file.
    Returns list of: ('heading', text), ('scene_break',), ('paragraph', text)
    """
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    elements = []
    in_metadata = False
    seen_headings = set()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if i == 0 and stripped == '---':
            i += 1
            while i < len(lines) and lines[i].strip() != '---':
                i += 1
            i += 1
            continue

        if stripped.startswith('# ') and 'CHAPTER' in stripped.upper():
            i += 1
            continue

        if is_metadata_line(stripped):
            in_metadata = True
            i += 1
            continue

        if stripped == '---':
            if in_metadata:
                in_metadata = False
                i += 1
                continue
            else:
                j = i + 1
                while j < len(lines) and lines[j].strip() == '':
                    j += 1
                if j < len(lines) and is_metadata_line(lines[j].strip()):
                    in_metadata = True
                    i += 1
                    continue
                if j < len(lines) and lines[j].strip().startswith('## '):
                    # Scene transition — emit scene_break
                    elements.append(('scene_break',))
                    i += 1
                    continue
                # Internal break within a scene — SKIP for clean output
                i += 1
                continue

        if in_metadata:
            if stripped == '':
                j = i + 1
                while j < len(lines) and lines[j].strip() == '':
                    j += 1
                if j < len(lines):
                    next_stripped = lines[j].strip()
                    if next_stripped.startswith('## ') and not is_metadata_line(next_stripped):
                        in_metadata = False
                    elif next_stripped == '---':
                        pass
                    elif not is_metadata_line(next_stripped) and not next_stripped.startswith('-') and not next_stripped.startswith('|') and not next_stripped.startswith('*'):
                        in_metadata = False
            i += 1
            continue

        if stripped == '':
            i += 1
            continue

        if stripped.startswith('|') and '|' in stripped[1:]:
            i += 1
            continue

        if stripped.startswith('- ') and any(kw in stripped for kw in [
            'threshold', 'PASS', 'adverb', 'metaphor', 'Kill', 'Remove',
            'Convert', 'Enhance', 'Strengthen', 'Adjust', 'Tighten',
            'Fix', 'Retain', 'Changed', 'Scene 1', 'Scene 2', 'Scene 3',
            'Sc1', 'Sc2', 'Sc3', 'Drift', 'Mass', 'Un-naming',
            'Cole', 'Chest', 'Root', 'Dad', 'Door', 'Postcard',
            'Hiss', 'Gate', 'Stairs', 'Map', 'Maya\'s',
            'Five-stage', 'Temperature', 'Cobblestones', 'Final line',
            'Somatic tell', 'NO copper', 'No jaw', 'Tell ARC',
            'No copper',
        ]):
            i += 1
            continue

        if stripped.startswith('## '):
            title = stripped[3:].strip()
            title = clean_scene_title(title)
            if title in seen_headings:
                i += 1
                continue
            seen_headings.add(title)
            elements.append(('heading', title))
            i += 1
            continue

        elements.append(('paragraph', stripped))
        i += 1

    # Post-process: ensure scene_break before every heading except the first
    fixed = []
    first_heading_seen = False
    for elem in elements:
        if elem[0] == 'heading':
            if first_heading_seen:
                if not fixed or fixed[-1][0] != 'scene_break':
                    fixed.append(('scene_break',))
            first_heading_seen = True
        fixed.append(elem)

    return fixed


def add_formatted_text(paragraph, text: str):
    """Parse markdown inline formatting and add runs to the paragraph."""
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|_(.+?)_)'

    pos = 0
    for match in re.finditer(pattern, text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])

        if match.group(2):  # ***bold italic***
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(3):  # **bold**
            run = paragraph.add_run(match.group(3))
            run.bold = True
        elif match.group(4):  # *italic*
            run = paragraph.add_run(match.group(4))
            run.italic = True
        elif match.group(5):  # _italic_
            run = paragraph.add_run(match.group(5))
            run.italic = True

        pos = match.end()

    if pos < len(text):
        paragraph.add_run(text[pos:])


def create_custom_style(doc, name, base_style='Normal'):
    """Create a custom paragraph style in the Word document."""
    styles = doc.styles
    style = styles.add_style(name, 1)  # 1 = WD_STYLE_TYPE.PARAGRAPH
    style.base_style = styles[base_style]
    return style


def build_docx():
    doc = Document()

    # Create custom Word styles that InDesign will map from
    style_chapter_num = create_custom_style(doc, 'ChapterNumber')
    style_chapter_title = create_custom_style(doc, 'ChapterTitle')
    style_scene_break = create_custom_style(doc, 'SceneBreak')

    for ch_idx, (rel_path, ch_number, ch_name) in enumerate(CHAPTERS):
        filepath = os.path.join(VAULT, rel_path)
        elements = extract_prose(filepath)

        # Strip leading scene breaks
        while elements and elements[0][0] == 'scene_break':
            elements.pop(0)

        # Add Chapter Number paragraph (just the digit)
        p = doc.add_paragraph()
        p.style = style_chapter_num
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(str(ch_number))

        # Add Chapter Title paragraph (chapter NAME from outline)
        p = doc.add_paragraph()
        p.style = style_chapter_title
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(ch_name)

        for elem in elements:
            if elem[0] == 'heading':
                # ALL scene headings are REMOVED for print
                continue
            elif elem[0] == 'scene_break':
                p = doc.add_paragraph()
                p.style = style_scene_break
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run('***')
            elif elem[0] == 'paragraph':
                p = doc.add_paragraph()
                p.style = doc.styles['Normal']
                add_formatted_text(p, elem[1])

        # Page break between chapters (except after last)
        if ch_idx < len(CHAPTERS) - 1:
            doc.add_page_break()

    doc.save(OUTPUT)
    print(f"Saved to: {OUTPUT}")
    print(f"Chapters: {len(CHAPTERS)}")

    # Stats
    total_words = 0
    chapter_nums = 0
    chapter_titles = 0
    scene_breaks = 0
    for para in doc.paragraphs:
        if para.style.name == 'ChapterNumber':
            chapter_nums += 1
        elif para.style.name == 'ChapterTitle':
            chapter_titles += 1
        elif para.style.name == 'SceneBreak':
            scene_breaks += 1
        elif para.style.name == 'Normal':
            total_words += len(para.text.split())

    print(f"Chapter numbers: {chapter_nums}")
    print(f"Chapter titles: {chapter_titles}")
    print(f"Scene breaks: {scene_breaks}")
    print(f"Prose words: {total_words}")


if __name__ == '__main__':
    build_docx()
